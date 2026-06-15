import base64
import json
import sqlite3
from collections.abc import Callable, Sequence
from io import BytesIO
from typing import Any, NotRequired, Protocol, TypedDict, cast
from uuid import uuid4

from docx import Document
from langgraph.graph import END, START, StateGraph

from citemind_worker.ark_gateway import ArkModelGateway
from citemind_worker.citation_validator import CitationValidator
from citemind_worker.model_catalog import (
    DEFAULT_ARK_BASE_URL,
    DEFAULT_CHAT_MODEL,
    DEFAULT_EMBEDDING_MODEL,
)
from citemind_worker.retrieval_service import HybridRetrievalService
from citemind_worker.storage import StorageRuntime

OUTLINE_SCHEMA: dict[str, object] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "summary": {"type": "string"},
        "sections": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "purpose": {"type": "string"},
                    "review_points": {"type": "array", "items": {"type": "string"}},
                    "evidence_chunk_ids": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["title", "purpose", "review_points", "evidence_chunk_ids"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["title", "summary", "sections"],
    "additionalProperties": False,
}

SECTION_SCHEMA: dict[str, object] = {
    "type": "object",
    "properties": {
        "paragraphs": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "text": {"type": "string"},
                    "evidence_chunk_ids": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["text", "evidence_chunk_ids"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["paragraphs"],
    "additionalProperties": False,
}


class WritingGateway(Protocol):
    async def generate_structured(
        self,
        request: dict[str, object],
        schema: dict[str, object],
    ) -> dict[str, object]: ...


class WorkflowState(TypedDict):
    project_id: str
    section_id: str
    knowledge_base_id: str
    goal: str
    workflow_type: str
    section_title: str
    section_purpose: str
    review_points: list[str]
    api_key: str
    base_url: str
    chat_model: str
    embedding_model: str
    retrieval: NotRequired[dict[str, object]]
    paragraphs: NotRequired[list[dict[str, object]]]
    audit: NotRequired[dict[str, object]]
    revision_feedback: NotRequired[dict[str, object]]


type GatewayFactory = Callable[[str, str, str], WritingGateway]


class WritingWorkflowService:
    def __init__(
        self,
        storage: StorageRuntime,
        *,
        retrieval: HybridRetrievalService | None = None,
        validator: CitationValidator | None = None,
        gateway_factory: GatewayFactory | None = None,
    ) -> None:
        self.storage = storage
        self.retrieval = retrieval or HybridRetrievalService(storage)
        self.validator = validator or CitationValidator(storage)
        self.gateway_factory = gateway_factory or _gateway_factory
        self.graph: Any = self._build_graph()

    def list_projects(self, knowledge_base_id: str) -> dict[str, object]:
        with self.storage.database.connect() as connection:
            exists = connection.execute(
                "SELECT 1 FROM knowledge_bases WHERE id = ?",
                (knowledge_base_id,),
            ).fetchone()
            if exists is None:
                raise ValueError("Knowledge base not found")
            rows = connection.execute(
                """
                SELECT id FROM writing_projects
                WHERE knowledge_base_id = ?
                ORDER BY updated_at DESC, created_at DESC
                """,
                (knowledge_base_id,),
            ).fetchall()
        return {
            "knowledgeBaseId": knowledge_base_id,
            "projects": [self._project_record(str(row["id"])) for row in rows],
        }

    def project(self, project_id: str) -> dict[str, object]:
        project = self._project_record(project_id)
        with self.storage.database.connect() as connection:
            section_rows = connection.execute(
                """
                SELECT id FROM writing_sections
                WHERE project_id = ?
                ORDER BY position
                """,
                (project_id,),
            ).fetchall()
            checkpoint_rows = connection.execute(
                """
                SELECT id, section_id, step, status, error_message, created_at
                FROM writing_checkpoints
                WHERE project_id = ?
                ORDER BY created_at DESC
                LIMIT 30
                """,
                (project_id,),
            ).fetchall()
        return {
            "project": project,
            "sections": [self._section_record(str(row["id"])) for row in section_rows],
            "checkpoints": [
                {
                    "id": str(row["id"]),
                    "sectionId": row["section_id"],
                    "step": str(row["step"]),
                    "status": str(row["status"]),
                    "errorMessage": row["error_message"],
                    "createdAt": str(row["created_at"]),
                }
                for row in checkpoint_rows
            ],
        }

    async def create_project(
        self,
        knowledge_base_id: str,
        *,
        goal: str,
        workflow_type: str,
        api_key: str,
        base_url: str = DEFAULT_ARK_BASE_URL,
        chat_model: str = DEFAULT_CHAT_MODEL,
        embedding_model: str = DEFAULT_EMBEDDING_MODEL,
    ) -> dict[str, object]:
        clean_goal = _required_text(goal, "写作目标")
        if workflow_type not in {"review", "article"}:
            raise ValueError("workflowType must be review or article")
        retrieval = await self.retrieval.retrieve(
            knowledge_base_id,
            clean_goal,
            api_key=api_key,
            base_url=base_url,
            embedding_model=embedding_model,
            limit=12,
            candidate_limit=36,
        )
        candidate_ids = _candidate_chunk_ids(retrieval)
        outline = await self.gateway_factory(api_key, base_url, chat_model).generate_structured(
            {
                "model": chat_model,
                "prompt": _outline_prompt(clean_goal, workflow_type, retrieval),
                "max_output_tokens": 1800,
            },
            OUTLINE_SCHEMA,
        )
        normalized_outline = _normalize_outline(outline, candidate_ids)
        project_id = f"writing-project-{uuid4().hex}"
        index_version_id = _index_version_id(retrieval)
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                INSERT INTO writing_projects(
                    id, knowledge_base_id, title, goal, workflow_type, status,
                    model_id, index_version_id, outline_json
                )
                VALUES (?, ?, ?, ?, ?, 'ready', ?, ?, ?)
                """,
                (
                    project_id,
                    knowledge_base_id,
                    normalized_outline["title"],
                    clean_goal,
                    workflow_type,
                    chat_model,
                    index_version_id,
                    _json(normalized_outline),
                ),
            )
            sections = cast(list[dict[str, object]], normalized_outline["sections"])
            for position, section in enumerate(sections):
                connection.execute(
                    """
                    INSERT INTO writing_sections(
                        id, project_id, position, title, purpose,
                        review_points_json, outline_evidence_json, status
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, 'pending')
                    """,
                    (
                        f"writing-section-{uuid4().hex}",
                        project_id,
                        position,
                        section["title"],
                        section["purpose"],
                        _json(section["reviewPoints"]),
                        _json(section["evidenceChunkIds"]),
                    ),
                )
            connection.commit()
        self._checkpoint(
            project_id,
            None,
            "outline",
            {"outline": normalized_outline, "indexVersionId": index_version_id},
        )
        return self.project(project_id)

    async def run_section(
        self,
        project_id: str,
        *,
        section_id: str | None,
        revise: bool,
        api_key: str,
        base_url: str = DEFAULT_ARK_BASE_URL,
        chat_model: str = DEFAULT_CHAT_MODEL,
        embedding_model: str = DEFAULT_EMBEDDING_MODEL,
    ) -> dict[str, object]:
        section = self._select_section(project_id, section_id)
        selected_id = str(section["id"])
        if revise:
            state = self._initial_state(
                project_id,
                section,
                api_key=api_key,
                base_url=base_url,
                chat_model=chat_model,
                embedding_model=embedding_model,
            )
            state["revision_feedback"] = _json_object(section["audit_json"])
        elif str(section["status"]) == "failed":
            state = self._resume_state(
                project_id,
                section,
                api_key=api_key,
                base_url=base_url,
                chat_model=chat_model,
                embedding_model=embedding_model,
            )
        else:
            state = self._initial_state(
                project_id,
                section,
                api_key=api_key,
                base_url=base_url,
                chat_model=chat_model,
                embedding_model=embedding_model,
            )
        self._mark_running(project_id, selected_id)
        try:
            await self.graph.ainvoke(state)
        except Exception as error:
            self._mark_failed(project_id, selected_id, str(error), state)
            raise ValueError(f"分节写作失败，可从检查点恢复：{error}") from error
        return self.project(project_id)

    def update_section(self, section_id: str, content: str) -> dict[str, object]:
        section = self._section_row(section_id)
        paragraphs = _paragraphs_from_content(content, _json_list(section["paragraphs_json"]))
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                UPDATE writing_sections
                SET content = ?, paragraphs_json = ?, audit_json = '{}',
                    status = 'needs_review', error_message = NULL,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (content.strip(), _json(paragraphs), section_id),
            )
            connection.execute("DELETE FROM writing_citations WHERE section_id = ?", (section_id,))
            connection.execute(
                """
                UPDATE writing_projects
                SET status = 'needs_revision',
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (section["project_id"],),
            )
            connection.commit()
        return self.project(str(section["project_id"]))

    def audit_section(self, section_id: str) -> dict[str, object]:
        section = self._section_row(section_id)
        project = self._project_row(str(section["project_id"]))
        paragraphs = _json_dict_list(section["paragraphs_json"])
        candidate_ids = list(
            dict.fromkeys(
                chunk_id
                for paragraph in paragraphs
                if isinstance(paragraph, dict)
                for chunk_id in _string_list(paragraph.get("evidenceChunkIds"))
            )
        )
        audit = self._audit(
            paragraphs,
            candidate_ids,
            str(project["index_version_id"] or ""),
        )
        self._persist_audit(str(project["id"]), section_id, paragraphs, audit)
        return self.project(str(project["id"]))

    def export_word(self, project_id: str) -> dict[str, object]:
        value = self.project(project_id)
        project = cast(dict[str, object], value["project"])
        sections = cast(list[dict[str, object]], value["sections"])
        document = Document()
        document.add_heading(str(project["title"]), level=0)
        document.add_paragraph(str(project["goal"]))
        for section in sections:
            document.add_heading(str(section["title"]), level=1)
            paragraphs = section.get("paragraphs")
            if isinstance(paragraphs, list):
                for paragraph in paragraphs:
                    if not isinstance(paragraph, dict):
                        continue
                    text = str(paragraph.get("text") or "")
                    evidence_ids = _string_list(paragraph.get("evidenceChunkIds"))
                    suffix = f" [{', '.join(evidence_ids)}]" if evidence_ids else ""
                    document.add_paragraph(f"{text}{suffix}")
        document.add_heading("引用证据", level=1)
        citations = [
            citation
            for section in sections
            for citation in cast(list[dict[str, object]], section.get("citations", []))
        ]
        if citations:
            for citation in citations:
                source = cast(dict[str, object], citation.get("source", {}))
                location = cast(dict[str, object], citation.get("location", {}))
                document.add_paragraph(
                    f"{citation.get('chunkId')} · {source.get('displayName')} · "
                    f"{_location_text(location)}"
                )
        else:
            document.add_paragraph("当前写作项目没有已验证引用。")
        output = BytesIO()
        document.save(output)
        return {
            "projectId": project_id,
            "fileName": f"{_safe_file_name(str(project['title']))}.docx",
            "base64": base64.b64encode(output.getvalue()).decode("ascii"),
        }

    def _build_graph(self) -> Any:
        graph = StateGraph(WorkflowState)
        graph.add_node("prepare", self._prepare)
        graph.add_node("retrieve", self._retrieve)
        graph.add_node("draft", self._draft)
        graph.add_node("audit", self._audit_node)
        graph.add_node("persist", self._persist)
        graph.add_edge(START, "prepare")
        graph.add_conditional_edges(
            "prepare",
            self._route,
            {"retrieve": "retrieve", "draft": "draft", "audit": "audit", "persist": "persist"},
        )
        graph.add_edge("retrieve", "draft")
        graph.add_edge("draft", "audit")
        graph.add_edge("audit", "persist")
        graph.add_edge("persist", END)
        return graph.compile()

    async def _prepare(self, state: WorkflowState) -> dict[str, object]:
        return {}

    @staticmethod
    def _route(state: WorkflowState) -> str:
        if "audit" in state:
            return "persist"
        if "paragraphs" in state:
            return "audit"
        if "retrieval" in state:
            return "draft"
        return "retrieve"

    async def _retrieve(self, state: WorkflowState) -> dict[str, object]:
        retrieval = await self.retrieval.retrieve(
            state["knowledge_base_id"],
            f"{state['section_title']} {state['section_purpose']}",
            api_key=state["api_key"],
            base_url=state["base_url"],
            embedding_model=state["embedding_model"],
            limit=12,
            candidate_limit=36,
        )
        patch: dict[str, object] = {"retrieval": retrieval}
        self._checkpoint(
            state["project_id"],
            state["section_id"],
            "retrieval",
            _merged_state(state, patch),
        )
        return patch

    async def _draft(self, state: WorkflowState) -> dict[str, object]:
        retrieval = state["retrieval"]
        candidate_ids = _candidate_chunk_ids(retrieval)
        draft = await self.gateway_factory(
            state["api_key"], state["base_url"], state["chat_model"]
        ).generate_structured(
            {
                "model": state["chat_model"],
                "prompt": _section_prompt(state, retrieval),
                "max_output_tokens": 2200,
            },
            SECTION_SCHEMA,
        )
        paragraphs = _normalize_paragraphs(draft.get("paragraphs"), candidate_ids)
        if not paragraphs:
            raise ValueError("模型未生成带有效候选引用的正文")
        patch: dict[str, object] = {"paragraphs": paragraphs}
        self._checkpoint(
            state["project_id"],
            state["section_id"],
            "draft",
            _merged_state(state, patch),
        )
        return patch

    async def _audit_node(self, state: WorkflowState) -> dict[str, object]:
        retrieval = state["retrieval"]
        audit = self._audit(
            state["paragraphs"],
            _candidate_chunk_ids(retrieval),
            _index_version_id(retrieval),
        )
        patch: dict[str, object] = {"audit": audit}
        self._checkpoint(
            state["project_id"],
            state["section_id"],
            "audit",
            _merged_state(state, patch),
        )
        return patch

    async def _persist(self, state: WorkflowState) -> dict[str, object]:
        self._persist_audit(
            state["project_id"],
            state["section_id"],
            state["paragraphs"],
            state["audit"],
        )
        self._checkpoint(state["project_id"], state["section_id"], "persist", state)
        return {}

    def _audit(
        self,
        paragraphs: Sequence[dict[str, object]],
        candidate_ids: Sequence[str],
        index_version_id: str,
    ) -> dict[str, object]:
        validator_paragraphs: list[dict[str, object]] = [
            {
                "text": str(paragraph.get("text") or ""),
                "evidence_chunk_ids": _string_list(paragraph.get("evidenceChunkIds")),
            }
            for paragraph in paragraphs
        ]
        validation = self.validator.validate(
            paragraphs=validator_paragraphs,
            candidate_chunk_ids=candidate_ids,
            index_version_id=index_version_id,
        )
        cited_ids = list(
            dict.fromkeys(
                chunk_id
                for paragraph in paragraphs
                for chunk_id in _string_list(paragraph.get("evidenceChunkIds"))
            )
        )
        conflicts = self._conflicts(cited_ids)
        invalid = cast(list[dict[str, object]], validation["invalidCitations"])
        suggestions = [
            {
                "type": "citation",
                "message": (
                    f"第 {_int_value(item.get('paragraphIndex')) + 1} 段需要补充或替换有效引用"
                ),
            }
            for item in invalid
        ]
        suggestions.extend(
            {
                "type": "conflict",
                "message": f"明确说明来源冲突：{item['sourceDisplayName']} 与 "
                f"{item['relatedDisplayName']}",
            }
            for item in conflicts
        )
        return {
            "valid": validation["valid"] is True and not conflicts,
            "citationValidation": validation,
            "conflicts": conflicts,
            "revisionSuggestions": suggestions,
        }

    def _conflicts(self, chunk_ids: Sequence[str]) -> list[dict[str, object]]:
        if not chunk_ids:
            return []
        placeholders = ",".join("?" for _ in chunk_ids)
        with self.storage.database.connect() as connection:
            source_rows = connection.execute(
                f"""
                SELECT DISTINCT s.id
                FROM chunks c
                JOIN source_versions sv ON sv.id = c.source_version_id
                JOIN sources s ON s.id = sv.source_id
                WHERE c.id IN ({placeholders})
                """,
                tuple(chunk_ids),
            ).fetchall()
            source_ids = [str(row["id"]) for row in source_rows]
            if not source_ids:
                return []
            source_placeholders = ",".join("?" for _ in source_ids)
            rows = connection.execute(
                f"""
                SELECT r.id, r.source_id, source.display_name AS source_name,
                       r.related_source_id, related.display_name AS related_name,
                       r.basis_json, r.confidence
                FROM source_relations r
                JOIN sources source ON source.id = r.source_id
                JOIN sources related ON related.id = r.related_source_id
                WHERE r.relation_type = 'conflicts'
                  AND r.status = 'confirmed'
                  AND r.source_id IN ({source_placeholders})
                  AND r.related_source_id IN ({source_placeholders})
                """,
                (*source_ids, *source_ids),
            ).fetchall()
        return [
            {
                "relationId": str(row["id"]),
                "sourceId": str(row["source_id"]),
                "sourceDisplayName": str(row["source_name"]),
                "relatedSourceId": str(row["related_source_id"]),
                "relatedDisplayName": str(row["related_name"]),
                "basis": _json_object(row["basis_json"]),
                "confidence": float(row["confidence"]),
            }
            for row in rows
        ]

    def _persist_audit(
        self,
        project_id: str,
        section_id: str,
        paragraphs: Sequence[dict[str, object]],
        audit: dict[str, object],
    ) -> None:
        status = "completed" if audit["valid"] is True else "needs_revision"
        content = "\n\n".join(str(paragraph.get("text") or "") for paragraph in paragraphs)
        validation = cast(dict[str, object], audit["citationValidation"])
        citations = cast(list[dict[str, object]], validation["validCitations"])
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                UPDATE writing_sections
                SET status = ?, content = ?, paragraphs_json = ?, audit_json = ?,
                    error_message = NULL,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (status, content, _json(paragraphs), _json(audit), section_id),
            )
            connection.execute("DELETE FROM writing_citations WHERE section_id = ?", (section_id,))
            for citation in citations:
                connection.execute(
                    """
                    INSERT OR IGNORE INTO writing_citations(
                        id, section_id, paragraph_index, chunk_id
                    )
                    VALUES (?, ?, ?, ?)
                    """,
                    (
                        f"writing-citation-{uuid4().hex}",
                        section_id,
                        citation["paragraphIndex"],
                        citation["chunkId"],
                    ),
                )
            remaining = connection.execute(
                """
                SELECT
                    SUM(CASE WHEN status = 'completed' THEN 1 ELSE 0 END) AS completed,
                    SUM(
                        CASE WHEN status IN ('needs_revision', 'failed') THEN 1 ELSE 0 END
                    ) AS blocked,
                    COUNT(*) AS total
                FROM writing_sections WHERE project_id = ?
                """,
                (project_id,),
            ).fetchone()
            project_status = "running"
            if int(remaining["blocked"] or 0) > 0:
                project_status = "needs_revision"
            elif int(remaining["completed"] or 0) == int(remaining["total"] or 0):
                project_status = "completed"
            connection.execute(
                """
                UPDATE writing_projects
                SET status = ?, error_message = NULL,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (project_status, project_id),
            )
            connection.commit()

    def _checkpoint(
        self,
        project_id: str,
        section_id: str | None,
        step: str,
        state: WorkflowState | dict[str, object],
        *,
        status: str = "completed",
        error_message: str | None = None,
    ) -> None:
        safe_state = {key: value for key, value in dict(state).items() if key != "api_key"}
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                INSERT INTO writing_checkpoints(
                    id, project_id, section_id, step, status, state_json, error_message
                )
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    f"writing-checkpoint-{uuid4().hex}",
                    project_id,
                    section_id,
                    step,
                    status,
                    _json(safe_state),
                    error_message,
                ),
            )
            connection.commit()

    def _resume_state(
        self,
        project_id: str,
        section: sqlite3.Row,
        *,
        api_key: str,
        base_url: str,
        chat_model: str,
        embedding_model: str,
    ) -> WorkflowState:
        with self.storage.database.connect() as connection:
            row = connection.execute(
                """
                SELECT state_json FROM writing_checkpoints
                WHERE section_id = ? AND status = 'completed'
                ORDER BY created_at DESC LIMIT 1
                """,
                (section["id"],),
            ).fetchone()
        if row is None:
            return self._initial_state(
                project_id,
                section,
                api_key=api_key,
                base_url=base_url,
                chat_model=chat_model,
                embedding_model=embedding_model,
            )
        values = _json_object(row["state_json"])
        values.update(
            {
                "api_key": api_key,
                "base_url": base_url,
                "chat_model": chat_model,
                "embedding_model": embedding_model,
            }
        )
        return cast(WorkflowState, values)

    def _initial_state(
        self,
        project_id: str,
        section: sqlite3.Row,
        *,
        api_key: str,
        base_url: str,
        chat_model: str,
        embedding_model: str,
    ) -> WorkflowState:
        project = self._project_row(project_id)
        return {
            "project_id": project_id,
            "section_id": str(section["id"]),
            "knowledge_base_id": str(project["knowledge_base_id"]),
            "goal": str(project["goal"]),
            "workflow_type": str(project["workflow_type"]),
            "section_title": str(section["title"]),
            "section_purpose": str(section["purpose"]),
            "review_points": _string_list(_json_value(section["review_points_json"])),
            "api_key": api_key,
            "base_url": base_url,
            "chat_model": chat_model,
            "embedding_model": embedding_model,
        }

    def _select_section(self, project_id: str, section_id: str | None) -> sqlite3.Row:
        self._project_row(project_id)
        with self.storage.database.connect() as connection:
            if section_id:
                row = connection.execute(
                    "SELECT * FROM writing_sections WHERE id = ? AND project_id = ?",
                    (section_id, project_id),
                ).fetchone()
            else:
                row = connection.execute(
                    """
                    SELECT * FROM writing_sections
                    WHERE project_id = ? AND status IN ('pending', 'failed')
                    ORDER BY position LIMIT 1
                    """,
                    (project_id,),
                ).fetchone()
        if row is None:
            raise ValueError("没有可继续写作的章节")
        return cast(sqlite3.Row, row)

    def _mark_running(self, project_id: str, section_id: str) -> None:
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                UPDATE writing_sections
                SET status = 'running', error_message = NULL,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (section_id,),
            )
            connection.execute(
                """
                UPDATE writing_projects
                SET status = 'running', error_message = NULL,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (project_id,),
            )
            connection.commit()

    def _mark_failed(
        self,
        project_id: str,
        section_id: str,
        message: str,
        state: WorkflowState,
    ) -> None:
        with self.storage.database.connect() as connection:
            connection.execute(
                """
                UPDATE writing_sections
                SET status = 'failed', error_message = ?, retry_count = retry_count + 1,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (message, section_id),
            )
            connection.execute(
                """
                UPDATE writing_projects
                SET status = 'failed', error_message = ?,
                    updated_at = strftime('%Y-%m-%dT%H:%M:%fZ', 'now')
                WHERE id = ?
                """,
                (message, project_id),
            )
            connection.commit()
        self._checkpoint(
            project_id,
            section_id,
            "workflow",
            state,
            status="failed",
            error_message=message,
        )

    def _project_record(self, project_id: str) -> dict[str, object]:
        row = self._project_row(project_id)
        return {
            "id": str(row["id"]),
            "knowledgeBaseId": str(row["knowledge_base_id"]),
            "title": str(row["title"]),
            "goal": str(row["goal"]),
            "workflowType": str(row["workflow_type"]),
            "status": str(row["status"]),
            "modelId": row["model_id"],
            "indexVersionId": row["index_version_id"],
            "outline": _json_object(row["outline_json"]),
            "audit": _json_object(row["audit_json"]),
            "errorMessage": row["error_message"],
            "createdAt": str(row["created_at"]),
            "updatedAt": str(row["updated_at"]),
        }

    def _section_record(self, section_id: str) -> dict[str, object]:
        row = self._section_row(section_id)
        with self.storage.database.connect() as connection:
            citations = connection.execute(
                """
                SELECT wc.paragraph_index, c.id AS chunk_id, c.page_number,
                       c.bounding_box_json, c.heading_path_json, c.anchor,
                       c.original_text, c.normalized_text,
                       s.id AS source_id, s.source_type, s.display_name, s.uri,
                       sv.id AS source_version_id
                FROM writing_citations wc
                JOIN chunks c ON c.id = wc.chunk_id
                JOIN source_versions sv ON sv.id = c.source_version_id
                JOIN sources s ON s.id = sv.source_id
                WHERE wc.section_id = ?
                ORDER BY wc.paragraph_index, wc.created_at
                """,
                (section_id,),
            ).fetchall()
        return {
            "id": str(row["id"]),
            "projectId": str(row["project_id"]),
            "position": int(row["position"]),
            "title": str(row["title"]),
            "purpose": str(row["purpose"]),
            "reviewPoints": _string_list(_json_value(row["review_points_json"])),
            "outlineEvidenceChunkIds": _string_list(_json_value(row["outline_evidence_json"])),
            "status": str(row["status"]),
            "content": str(row["content"]),
            "paragraphs": _json_list(row["paragraphs_json"]),
            "audit": _json_object(row["audit_json"]),
            "citations": [_citation_record(citation) for citation in citations],
            "errorMessage": row["error_message"],
            "retryCount": int(row["retry_count"]),
            "createdAt": str(row["created_at"]),
            "updatedAt": str(row["updated_at"]),
        }

    def _project_row(self, project_id: str) -> sqlite3.Row:
        with self.storage.database.connect() as connection:
            row = connection.execute(
                "SELECT * FROM writing_projects WHERE id = ?",
                (project_id,),
            ).fetchone()
        if row is None:
            raise ValueError("Writing project not found")
        return cast(sqlite3.Row, row)

    def _section_row(self, section_id: str) -> sqlite3.Row:
        with self.storage.database.connect() as connection:
            row = connection.execute(
                "SELECT * FROM writing_sections WHERE id = ?",
                (section_id,),
            ).fetchone()
        if row is None:
            raise ValueError("Writing section not found")
        return cast(sqlite3.Row, row)


def _gateway_factory(api_key: str, base_url: str, _chat_model: str) -> WritingGateway:
    return ArkModelGateway(api_key, base_url=base_url)


def _outline_prompt(goal: str, workflow_type: str, retrieval: dict[str, object]) -> str:
    mode = "知识点总结与复习提纲" if workflow_type == "review" else "基于证据的写作大纲"
    return (
        f"任务：为用户目标生成{mode}。\n目标：{goal}\n"
        "每一节必须绑定下列检索证据中的 chunk_id，不得编造引用。"
        "复习提纲应突出关键知识点、易错点与复习问题；写作大纲应说明每节目的和论证重点。\n"
        f"证据：{_json(retrieval.get('context', []))}"
    )


def _section_prompt(state: WorkflowState, retrieval: dict[str, object]) -> str:
    feedback = state.get("revision_feedback")
    return (
        f"围绕写作目标“{state['goal']}”撰写章节“{state['section_title']}”。\n"
        f"章节目的：{state['section_purpose']}\n"
        f"复习或论证要点：{_json(state['review_points'])}\n"
        "每个事实性段落必须列出 evidence_chunk_ids，且只能使用证据中的 chunk_id。"
        "若证据存在冲突，应在正文中明确说明，不得静默选择一方。\n"
        f"上次审计与修订建议：{_json(feedback or {})}\n"
        f"证据：{_json(retrieval.get('context', []))}"
    )


def _normalize_outline(value: dict[str, object], candidate_ids: Sequence[str]) -> dict[str, object]:
    candidate_set = set(candidate_ids)
    sections: list[dict[str, object]] = []
    raw_sections = value.get("sections")
    if isinstance(raw_sections, list):
        for raw in raw_sections[:12]:
            if not isinstance(raw, dict):
                continue
            title = _required_text(str(raw.get("title") or ""), "章节标题")
            purpose = _required_text(str(raw.get("purpose") or ""), "章节目的")
            evidence_ids = [
                chunk_id
                for chunk_id in _string_list(raw.get("evidence_chunk_ids"))
                if chunk_id in candidate_set
            ]
            sections.append(
                {
                    "title": title,
                    "purpose": purpose,
                    "reviewPoints": _string_list(raw.get("review_points")),
                    "evidenceChunkIds": evidence_ids,
                }
            )
    if not sections:
        raise ValueError("模型未生成有效写作大纲")
    return {
        "title": _required_text(str(value.get("title") or ""), "项目标题"),
        "summary": str(value.get("summary") or "").strip(),
        "sections": sections,
    }


def _normalize_paragraphs(value: object, candidate_ids: Sequence[str]) -> list[dict[str, object]]:
    if not isinstance(value, list):
        return []
    candidate_set = set(candidate_ids)
    paragraphs: list[dict[str, object]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text") or "").strip()
        if not text:
            continue
        evidence_ids = [
            chunk_id
            for chunk_id in _string_list(item.get("evidence_chunk_ids"))
            if chunk_id in candidate_set
        ]
        paragraphs.append({"text": text, "evidenceChunkIds": evidence_ids})
    return paragraphs


def _paragraphs_from_content(
    content: str,
    previous: Sequence[object],
) -> list[dict[str, object]]:
    previous_paragraphs = [item for item in previous if isinstance(item, dict)]
    blocks = [block.strip() for block in content.split("\n\n") if block.strip()]
    return [
        {
            "text": block,
            "evidenceChunkIds": _string_list(
                previous_paragraphs[index].get("evidenceChunkIds")
                if index < len(previous_paragraphs)
                else None
            ),
        }
        for index, block in enumerate(blocks)
    ]


def _candidate_chunk_ids(retrieval: dict[str, object]) -> list[str]:
    results = retrieval.get("results")
    if not isinstance(results, list):
        return []
    return [
        str(item["chunkId"])
        for item in results
        if isinstance(item, dict) and isinstance(item.get("chunkId"), str)
    ]


def _index_version_id(retrieval: dict[str, object]) -> str:
    index = retrieval.get("indexVersion")
    return str(index.get("id") or "") if isinstance(index, dict) else ""


def _citation_record(row: sqlite3.Row) -> dict[str, object]:
    return {
        "paragraphIndex": int(row["paragraph_index"]),
        "chunkId": str(row["chunk_id"]),
        "source": {
            "id": str(row["source_id"]),
            "versionId": str(row["source_version_id"]),
            "type": str(row["source_type"]),
            "displayName": str(row["display_name"]),
            "uri": row["uri"],
        },
        "location": {
            "pageNumber": row["page_number"],
            "boundingBox": _json_object(row["bounding_box_json"]),
            "headingPath": _string_list(_json_value(row["heading_path_json"])),
            "anchor": row["anchor"],
        },
        "text": {
            "original": str(row["original_text"]),
            "normalized": str(row["normalized_text"]),
            "preview": str(row["normalized_text"])[:240],
        },
    }


def _location_text(location: dict[str, object]) -> str:
    if location.get("pageNumber"):
        return f"第 {location['pageNumber']} 页"
    headings = _string_list(location.get("headingPath"))
    if headings:
        return " / ".join(headings)
    return str(location.get("anchor") or "位置未标注")


def _required_text(value: str, label: str) -> str:
    clean = " ".join(value.split())
    if not clean:
        raise ValueError(f"{label}不能为空")
    return clean


def _safe_file_name(value: str) -> str:
    clean = "".join(character if character not in '\\/:*?"<>|' else "-" for character in value)
    return clean.strip()[:80] or "citeMind-writing"


def _json(value: object) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _json_value(value: object) -> object:
    if not isinstance(value, str) or not value:
        return None
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return None


def _json_object(value: object) -> dict[str, object]:
    parsed = _json_value(value)
    return parsed if isinstance(parsed, dict) else {}


def _json_list(value: object) -> list[object]:
    parsed = _json_value(value)
    return parsed if isinstance(parsed, list) else []


def _json_dict_list(value: object) -> list[dict[str, object]]:
    return [item for item in _json_list(value) if isinstance(item, dict)]


def _string_list(value: object) -> list[str]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, str) and item]


def _merged_state(state: WorkflowState, patch: dict[str, object]) -> dict[str, object]:
    values: dict[str, object] = dict(state)
    values.update(patch)
    return values


def _int_value(value: object) -> int:
    return value if isinstance(value, int) else 0
